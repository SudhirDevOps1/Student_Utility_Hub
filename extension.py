import streamlit as st
from PIL import Image
import io
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import PyPDF2
import pdfplumber
from docx import Document
import base64

def show():
    st.markdown("## 🧩 File Converter Engine")
    
    tabs = st.tabs(["Text & Web", "Office Files", "Image & Media", "PDF Tools"])
    
    # TEXT & WEB
    with tabs[0]:
        st.markdown("### 📝 Text & Web Converters")
        
        converter = st.selectbox("Select Converter", [
            "HTML → Text",
            "HTML → PDF",
            "Text → Image",
            "Markdown → PDF"
        ])
        
        if converter == "HTML → Text":
            html_input = st.text_area("Paste HTML:", height=200)
            if st.button("Convert"):
                from html.parser import HTMLParser
                class HTMLTextExtractor(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.text = []
                    def handle_data(self, data):
                        self.text.append(data)
                
                parser = HTMLTextExtractor()
                parser.feed(html_input)
                text = ' '.join(parser.text)
                st.text_area("Extracted Text:", text, height=200)
                st.download_button("Download", text, "output.txt", "text/plain")
        
        elif converter == "HTML → PDF":
            html_input = st.text_area("Paste HTML:", height=200)
            if st.button("Convert to PDF"):
                buffer = io.BytesIO()
                c = canvas.Canvas(buffer, pagesize=letter)
                c.drawString(50, 750, "HTML to PDF Conversion")
                c.drawString(50, 700, html_input[:100])
                c.save()
                buffer.seek(0)
                st.download_button("Download PDF", buffer, "output.pdf", "application/pdf")
        
        elif converter == "Text → Image":
            text_input = st.text_area("Enter text:", height=200)
            if st.button("Convert to Image"):
                from PIL import ImageDraw, ImageFont
                img = Image.new('RGB', (800, 600), color='white')
                draw = ImageDraw.Draw(img)
                draw.text((50, 50), text_input, fill='black')
                
                buffer = io.BytesIO()
                img.save(buffer, format='PNG')
                buffer.seek(0)
                
                st.image(buffer, caption="Generated Image")
                st.download_button("Download", buffer, "text_image.png", "image/png")
        
        elif converter == "Markdown → PDF":
            md_input = st.text_area("Enter Markdown:", height=200)
            if st.button("Convert to PDF"):
                buffer = io.BytesIO()
                c = canvas.Canvas(buffer, pagesize=letter)
                c.drawString(50, 750, "Markdown to PDF")
                lines = md_input.split('\n')
                y = 700
                for line in lines[:30]:
                    c.drawString(50, y, line)
                    y -= 20
                c.save()
                buffer.seek(0)
                st.download_button("Download PDF", buffer, "markdown.pdf", "application/pdf")
    
    # OFFICE FILES
    with tabs[1]:
        st.markdown("### 📊 Office File Converters")
        
        converter = st.selectbox("Select Converter", [
            "DOCX → Text",
            "DOCX → PDF",
            "CSV → Excel",
            "Excel → CSV"
        ], key="office")
        
        if converter == "DOCX → Text":
            uploaded = st.file_uploader("Upload DOCX", type=['docx'])
            if uploaded and st.button("Extract Text"):
                doc = Document(uploaded)
                text = '\n'.join([para.text for para in doc.paragraphs])
                st.text_area("Extracted Text:", text, height=300)
                st.download_button("Download", text, "output.txt", "text/plain")
        
        elif converter == "DOCX → PDF":
            uploaded = st.file_uploader("Upload DOCX", type=['docx'])
            if uploaded and st.button("Convert to PDF"):
                doc = Document(uploaded)
                
                buffer = io.BytesIO()
                c = canvas.Canvas(buffer, pagesize=letter)
                y = 750
                for para in doc.paragraphs[:30]:
                    c.drawString(50, y, para.text[:80])
                    y -= 20
                    if y < 50:
                        break
                c.save()
                buffer.seek(0)
                st.download_button("Download PDF", buffer, "output.pdf", "application/pdf")
        
        elif converter == "CSV → Excel":
            uploaded = st.file_uploader("Upload CSV", type=['csv'])
            if uploaded and st.button("Convert"):
                df = pd.read_csv(uploaded)
                st.dataframe(df)
                
                buffer = io.BytesIO()
                with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False)
                buffer.seek(0)
                st.download_button("Download Excel", buffer, "output.xlsx", 
                                 "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        
        elif converter == "Excel → CSV":
            uploaded = st.file_uploader("Upload Excel", type=['xlsx', 'xls'])
            if uploaded and st.button("Convert"):
                df = pd.read_excel(uploaded)
                st.dataframe(df)
                
                csv = df.to_csv(index=False)
                st.download_button("Download CSV", csv, "output.csv", "text/csv")
    
    # IMAGE & MEDIA
    with tabs[2]:
        st.markdown("### 🖼️ Image & Media Converters")
        
        converter = st.selectbox("Select Converter", [
            "Image → PDF",
            "Image → Text (OCR)",
            "Video → Audio (MP3)"
        ], key="media")
        
        if converter == "Image → PDF":
            uploaded = st.file_uploader("Upload Image", type=['png', 'jpg', 'jpeg'])
            if uploaded and st.button("Convert"):
                image = Image.open(uploaded)
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                buffer = io.BytesIO()
                image.save(buffer, format='PDF')
                buffer.seek(0)
                st.download_button("Download PDF", buffer, "image.pdf", "application/pdf")
        
        elif converter == "Image → Text (OCR)":
            st.info("Use the OCR & Text Tools section for advanced OCR features")
            uploaded = st.file_uploader("Upload Image", type=['png', 'jpg', 'jpeg'])
            if uploaded and st.button("Extract Text"):
                try:
                    import pytesseract
                    image = Image.open(uploaded)
                    text = pytesseract.image_to_string(image)
                    st.text_area("Extracted Text:", text, height=300)
                except:
                    st.error("Tesseract not installed")
        
        elif converter == "Video → Audio (MP3)":
            st.warning("This feature requires moviepy library")
            uploaded = st.file_uploader("Upload Video", type=['mp4', 'avi', 'mov'])
            if uploaded:
                st.info("Video to audio conversion requires moviepy. Install: pip install moviepy")
    
    # PDF TOOLS
    with tabs[3]:
        st.markdown("### 📄 PDF Converters")
        
        converter = st.selectbox("Select Converter", [
            "PDF → Text",
            "PDF → Images"
        ], key="pdf_conv")
        
        if converter == "PDF → Text":
            uploaded = st.file_uploader("Upload PDF", type=['pdf'])
            if uploaded and st.button("Extract Text"):
                with pdfplumber.open(uploaded) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() + "\n\n"
                
                st.text_area("Extracted Text:", text, height=300)
                st.download_button("Download", text, "pdf_text.txt", "text/plain")
        
        elif converter == "PDF → Images":
            uploaded = st.file_uploader("Upload PDF", type=['pdf'])
            if uploaded:
                st.info("PDF to images requires PyMuPDF. Install: pip install PyMuPDF")
                if st.button("Convert"):
                    try:
                        import fitz
                        doc = fitz.open(stream=uploaded.read(), filetype="pdf")
                        
                        for page_num in range(min(5, len(doc))):
                            page = doc[page_num]
                            pix = page.get_pixmap(dpi=150)
                            img_data = pix.tobytes("png")
                            
                            st.image(img_data, caption=f"Page {page_num + 1}")
                            st.download_button(f"Download Page {page_num + 1}", 
                                             img_data, f"page_{page_num + 1}.png", "image/png",
                                             key=f"dl_{page_num}")
                    except:
                        st.error("PyMuPDF not available")
    
    st.markdown("---")
    st.info("💡 Tip: For more advanced features, check other tool sections")
